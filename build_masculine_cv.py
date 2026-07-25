import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="none"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def generate_masculine_cv():
    doc = Document()

    # Margins (0.3 inch)
    for s in doc.sections:
        s.top_margin = Inches(0.3)
        s.bottom_margin = Inches(0.3)
        s.left_margin = Inches(0.3)
        s.right_margin = Inches(0.3)

    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Segoe UI'
    style_normal.font.size = Pt(9)

    # EXECUTIVE TECH MASCULINE PALETTE
    SIDEBAR_FILL = "0F172A"       # Deep Slate Navy
    MAIN_FILL = "FFFFFF"          # Pure Crisp White
    
    CYAN_BLUE_HEADER = RGBColor(0x38, 0xBD, 0xF8)   # Electric Sky Blue
    ICE_WHITE_TEXT = RGBColor(0xF8, 0xFA, 0xFC)     # Crisp White
    SLATE_LIGHT = RGBColor(0xCB, 0xD5, 0xE1)        # Light Slate Gray
    
    NAVY_TITLE = RGBColor(0x0A, 0x11, 0x28)         # Deep Executive Navy
    OCEAN_BLUE = RGBColor(0x02, 0x84, 0xC7)         # Professional Ocean Blue Accent
    BODY_DARK = RGBColor(0x33, 0x41, 0x55)          # Slate Text Body
    SUBTLE_TEXT = RGBColor(0x64, 0x74, 0x8B)        # Muted Slate

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)

    col_widths = [Inches(2.5), Inches(4.9)]

    # ==================== PAGE 1 / ROW 0 CELL 0 (SIDEBAR TOP) ====================
    c00 = table.cell(0, 0)
    c00.width = col_widths[0]
    set_cell_background(c00, SIDEBAR_FILL)
    set_cell_margins(c00, top=180, bottom=180, left=180, right=180)

    # --- PHOTO FRAME PLACEHOLDER (Top Left) ---
    photo_box = c00.add_table(rows=1, cols=1)
    photo_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_cell = photo_box.cell(0, 0)
    p_cell.width = Inches(1.8)
    set_cell_background(p_cell, "1E293B")
    set_cell_margins(p_cell, top=200, bottom=200, left=100, right=100)
    
    p_ph = p_cell.paragraphs[0]
    p_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ph = p_ph.add_run("📷 EMPLACEMENT PHOTO CV\n(Insérer votre photo ici)")
    r_ph.font.size = Pt(7.5)
    r_ph.font.bold = True
    r_ph.font.color.rgb = RGBColor(0x38, 0xBD, 0xF8)

    c00.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_sb_h(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text.upper())
        r.font.name = 'Segoe UI Semibold'
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = CYAN_BLUE_HEADER

    def add_sb_t(cell, icon_txt):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(icon_txt)
        r.font.size = Pt(8.5)
        r.font.color.rgb = ICE_WHITE_TEXT

    add_sb_h(c00, "CONTACT")
    add_sb_t(c00, "✉  koagervais85@gmail.com")
    add_sb_t(c00, "✆  +237 695 35 34 02")
    add_sb_t(c00, "⌂  Douala / Ngaoundéré, Cameroun")
    add_sb_t(c00, "🌐  github.com/gervais-afk")

    add_sb_h(c00, "COMPÉTENCES & OUTILS")
    add_sb_t(c00, "Google Gemma 4 / Gemini   ■ ■ ■ ■ ■")
    add_sb_t(c00, "Neo4j GraphRAG (N10S)    ■ ■ ■ ■ ■")
    add_sb_t(c00, "Firebase Genkit / RAG     ■ ■ ■ ■ ■")
    add_sb_t(c00, "Python (BAEL 91 Sandbox)  ■ ■ ■ ■ ■")
    add_sb_t(c00, "IfcOpenShell (5D BIM)     ■ ■ ■ ■ □")
    add_sb_t(c00, "PostgreSQL / SQL         ■ ■ ■ ■ ■")
    add_sb_t(c00, "Next.js 14 / React        ■ ■ ■ ■ □")
    add_sb_t(c00, "FastAPI / Node.js         ■ ■ ■ ■ □")
    add_sb_t(c00, "Streamlit / PWA ML        ■ ■ ■ ■ ■")
    add_sb_t(c00, "MLflow (MLOps) / Docker   ■ ■ ■ ■ □")

    add_sb_h(c00, "LANGUES")
    add_sb_t(c00, "Français   ■ ■ ■ ■ ■  (Courant)")
    add_sb_t(c00, "Anglais    ■ ■ ■ ■ □  (Pro / Tech)")

    add_sb_h(c00, "ATOUTS CLÉS")
    add_sb_t(c00, "◈  Double compétence IA & Génie Civil")
    add_sb_t(c00, "◈  Gestion des risques & Sécurité (AVSEC)")
    add_sb_t(c00, "◈  Guardrails IA & Calculs déterministes")
    add_sb_t(c00, "◈  Rigueur & Souveraineté logicielle")

    # ==================== PAGE 1 / ROW 0 CELL 1 (MAIN COLUMN TOP) ====================
    c01 = table.cell(0, 1)
    c01.width = col_widths[1]
    set_cell_background(c01, MAIN_FILL)
    set_cell_margins(c01, top=180, bottom=180, left=200, right=180)

    p_m01 = c01.paragraphs[0]
    p_m01.paragraph_format.space_before = Pt(0)
    p_m01.paragraph_format.space_after = Pt(2)
    r_nm = p_m01.add_run("KOA MARIE GERVAIS NELLY")
    r_nm.font.name = 'Segoe UI'
    r_nm.font.bold = True
    r_nm.font.size = Pt(18)
    r_nm.font.color.rgb = NAVY_TITLE

    p_sub = c01.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(3)
    r_sb = p_sub.add_run("Lead AI Engineer & Consultant IA / Data   │   Fondateur @ Archi Cam AI")
    r_sb.font.size = Pt(10)
    r_sb.font.bold = True
    r_sb.font.color.rgb = OCEAN_BLUE

    p_dec = c01.add_paragraph()
    p_dec.paragraph_format.space_before = Pt(0)
    p_dec.paragraph_format.space_after = Pt(6)
    r_d = p_dec.add_run("─── ◈ ────────────────────────────────────────────────── ◈ ───")
    r_d.font.size = Pt(8)
    r_d.font.color.rgb = OCEAN_BLUE

    def add_mn_h(cell, title):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run("◈  ")
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = OCEAN_BLUE
        r2 = p.add_run(title.upper())
        r2.font.name = 'Segoe UI'
        r2.font.bold = True
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = NAVY_TITLE

    add_mn_h(c01, "PROFIL SYNTHÉTIQUE")
    p_pr = c01.add_paragraph()
    p_pr.paragraph_format.space_before = Pt(0)
    p_pr.paragraph_format.space_after = Pt(6)
    p_pr.paragraph_format.line_spacing = 1.15
    r_pr = p_pr.add_run(
        "Ingénieur & Spécialiste en IA Appliquée alliant la rigueur du Génie Civil et l'expertise en Sûreté Opérationnelle à la maîtrise des architectures émergentes d'IA Générative, GraphRAG (Neo4j) et MLOps. Fondateur d'Archi Cam AI (projet développé pour le Google Africa Applied AI Lab). Concepteur de solutions IA souveraines d'entreprise éliminant les hallucinations des LLM par l'intégration de moteurs déterministes en Sandbox."
    )
    r_pr.font.size = Pt(9)
    r_pr.font.color.rgb = BODY_DARK

    add_mn_h(c01, "PROJETS IA MAJEURS & RÉALISATIONS")

    def add_proj(cell, name, sub_badge, bullets):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(name + "  ")
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = NAVY_TITLE
        
        r2 = p.add_run("–  " + sub_badge)
        r2.font.italic = True
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = OCEAN_BLUE

        for b in bullets:
            bp = cell.add_paragraph()
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(1.5)
            bp.paragraph_format.left_indent = Inches(0.12)
            rb_ico = bp.add_run("▸  ")
            rb_ico.font.bold = True
            rb_ico.font.color.rgb = OCEAN_BLUE
            rb_txt = bp.add_run(b)
            rb_txt.font.size = Pt(8.5)
            rb_txt.font.color.rgb = BODY_DARK

    add_proj(
        c01,
        "Archi Cam AI 🏛️",
        "SaaS IA Agentique & 5D BIM (Candidature Google Africa Applied AI Lab)",
        [
            "Plateforme IA souveraine de modélisation BIM 5D et génération automatisée de devis BTP normés (BAEL 91 / Eurocodes).",
            "Combinaison de Google Gemma 4 12B local, Gemini 1.5 Pro Vision et d'un moteur Python Sandbox (IfcOpenShell, BAEL 91).",
            "Génération de devis Excel (DQE) en < 2 min et rendus HD via Imagen 3 + ControlNet."
        ]
    )

    add_proj(
        c01,
        "Sovereign.BI Agentic 📊",
        "Business Intelligence Agentique & Guardrails",
        [
            "Moteur décisionnel autonome d'interrogation de bases SQL complexes en langage naturel (React, FastAPI, PostgreSQL).",
            "Architecture TypeScript Orchestrator, Neo4j N10S (GraphRAG) et auditeur d'explicabilité SHAP Sentinel."
        ]
    )

    add_proj(
        c01,
        "Dataset Automator ⚙️ & VigieSahel 🌾",
        "MLOps & IA Impact Climat / Santé",
        [
            "Dataset Automator : Pipeline RAG d'évaluation de séries temporelles (Neo4j, MLflow, Genkit, Gemma-2).",
            "VigieSahel : Plateforme prédictive des semis agricoles et suivi PM2.5 / méningite (Streamlit, Supabase, ML)."
        ]
    )

    # ==================== PAGE 2 / ROW 1 CELL 0 (SIDEBAR BOTTOM) ====================
    c10 = table.cell(1, 0)
    c10.width = col_widths[0]
    set_cell_background(c10, SIDEBAR_FILL)
    set_cell_margins(c10, top=140, bottom=180, left=180, right=180)

    p_r10 = c10.paragraphs[0]
    p_r10.paragraph_format.space_before = Pt(0)
    p_r10.paragraph_format.space_after = Pt(2)
    r_sb_nm = p_r10.add_run("KOA MARIE GERVAIS NELLY")
    r_sb_nm.font.bold = True
    r_sb_nm.font.size = Pt(9)
    r_sb_nm.font.color.rgb = CYAN_BLUE_HEADER

    p_sb_dec = c10.add_paragraph()
    p_sb_dec.paragraph_format.space_before = Pt(0)
    p_sb_dec.paragraph_format.space_after = Pt(2)
    r_sbd = p_sb_dec.add_run("── ◈ ──")
    r_sbd.font.size = Pt(8)
    r_sbd.font.color.rgb = CYAN_BLUE_HEADER

    add_sb_h(c10, "DISPONIBILITÉ & ENGAGEMENT")
    add_sb_t(c10, "◦  Statut : Disponible / Conseil")
    add_sb_t(c10, "◦  Mobilité : Internationale / Remote")

    add_sb_h(c10, "RÉFÉRENCES")
    add_sb_t(c10, "◈  Disponibles sur demande")
    add_sb_t(c10, "◈ ◈ ◈")

    # ==================== PAGE 2 / ROW 1 CELL 1 (MAIN COLUMN BOTTOM) ====================
    c11 = table.cell(1, 1)
    c11.width = col_widths[1]
    set_cell_background(c11, MAIN_FILL)
    set_cell_margins(c11, top=140, bottom=180, left=200, right=180)

    p_r11 = c11.paragraphs[0]
    p_r11.paragraph_format.space_before = Pt(0)
    p_r11.paragraph_format.space_after = Pt(2)

    add_mn_h(c11, "PARCOURS PROFESSIONNEL")

    def add_job(cell, title, period, company, bullets):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(title + "  ")
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = NAVY_TITLE

        r2 = p.add_run(f"  {period}")
        r2.font.italic = True
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = OCEAN_BLUE

        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        r3 = p2.add_run(company)
        r3.font.bold = True
        r3.font.size = Pt(8.5)
        r3.font.color.rgb = SUBTLE_TEXT

        for b in bullets:
            bp = cell.add_paragraph()
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(1)
            bp.paragraph_format.left_indent = Inches(0.12)
            rb_ico = bp.add_run("▸  ")
            rb_ico.font.bold = True
            rb_ico.font.color.rgb = OCEAN_BLUE
            rb_txt = bp.add_run(b)
            rb_txt.font.size = Pt(8.5)
            rb_txt.font.color.rgb = BODY_DARK

    add_job(
        c11,
        "Consultant IA & Data Science",
        "2025 – Présent",
        "Projets Indépendants & Entreprises   │   Douala",
        [
            "Analyse exploratoire et prétraitement de jeux de données massifs complexes.",
            "Modélisation de graphes de connaissances (Neo4j Cypher) et développement de pipelines RAG.",
            "Conception de bases de données SQL/PostgreSQL et reporting décisionnel interactif."
        ]
    )

    add_job(
        c11,
        "Agent de Sûreté Aéroportuaire (AVSEC)",
        "2018 – Présent",
        "CCAA (Autorité Aéronautique du Cameroun)",
        [
            "Analyse des risques opérationnels critiques et contrôle strict des accès sécurisés.",
            "Rédaction de rapports d'audit de sûreté et coordination d'interventions opérationnelles."
        ]
    )

    add_mn_h(c11, "FORMATION ACADÉMIQUE")

    def add_edu(cell, degree, period, school, details):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(degree + "  ")
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = NAVY_TITLE

        r2 = p.add_run(f"  {period}")
        r2.font.italic = True
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = OCEAN_BLUE

        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        r3 = p2.add_run(school + "   │   " + details)
        r3.font.size = Pt(8.5)
        r3.font.color.rgb = SUBTLE_TEXT

    add_edu(c11, "Master 2 Intelligence Artificielle & Data Science", "2025 – 2027", "Université de Ngaoundéré", "Spécialisation Graphes (Neo4j), MLOps & LLM")
    add_edu(c11, "Licence & BTS Génie Civil (Option Bâtiment)", "2015 – 2016", "ISTDI / IUC Douala", "Dimensionnement BAEL 91 & Métrés BTP")

    # Bottom Quote
    p_q = c11.add_paragraph()
    p_q.paragraph_format.space_before = Pt(8)
    p_q.paragraph_format.space_after = Pt(0)
    r_q = p_q.add_run("« Motivé, rigoureux et engagé à concevoir des solutions IA souveraines et fiables. »")
    r_q.font.italic = True
    r_q.font.size = Pt(8)
    r_q.font.color.rgb = SUBTLE_TEXT

    # Output paths
    f1 = r"c:\Users\HP\Desktop\portfolio-gervais\KOA_MARIE_GERVAIS_NELLY_CV_EXECUTIVE.docx"
    f2 = r"c:\Users\HP\Desktop\portfolio-gervais\KOA_MARIE_GERVAIS_NELLY_CV.docx"
    
    try:
        doc.save(f1)
    except Exception as e:
        print(f"Note on f1: {e}")
        
    try:
        doc.save(f2)
    except Exception as e:
        print(f"Note on f2: {e}")

    print(f"Generated masculine executive CV successfully at:\n  {f1}\n  {f2}")

if __name__ == "__main__":
    generate_masculine_cv()
