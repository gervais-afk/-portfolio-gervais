import docx

doc = docx.Document('CV_TELLENG ALBA.docx')

print("=== DOCX DETAILED PARSING ===")
for ti, t in enumerate(doc.tables):
    for ri, r in enumerate(t.rows):
        for ci, c in enumerate(r.cells):
            tcPr = c._tc.get_or_add_tcPr()
            shd = tcPr.find(docx.oxml.ns.qn('w:shd'))
            fill = shd.get(docx.oxml.ns.qn('w:fill')) if shd is not None else 'None'
            print(f"\n--- Cell [{ri},{ci}] (Width: {c.width}, Fill: {fill}) ---")
            for p in c.paragraphs:
                txt = p.text.strip()
                if txt:
                    colors = []
                    fonts = []
                    for run in p.runs:
                        if run.font.color and run.font.color.rgb:
                            colors.append(str(run.font.color.rgb))
                        if run.font.name:
                            fonts.append(run.font.name)
                    print(f"P: '{txt}' | Fonts: {set(fonts)} | Colors: {set(colors)}")
