import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_bottom_border(paragraph, color_hex="00A8B5", size="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_cv_docx():
    doc = Document()
    
    # Page Margins (0.5 inch / 1.27 cm)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.header_distance = Inches(0.2)
        section.footer_distance = Inches(0.2)
        
    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Segoe UI'
    style_normal.font.size = Pt(9.5)
    style_normal.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    # Palette
    NAVY = RGBColor(0x0A, 0x11, 0x28)
    CYAN = RGBColor(0x00, 0x8B, 0x8B)
    DARK_BLUE = RGBColor(0x0D, 0x1B, 0x3E)
    GRAY_TEXT = RGBColor(0x4A, 0x55, 0x68)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    # 1. HEADER BANNER TABLE
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = header_table.cell(0, 0)
    cell.width = Inches(7.5)
    set_cell_background(cell, "0A1128")
    set_cell_margins(cell, top=200, bottom=200, left=240, right=240)

    # Header Name
    p_name = cell.paragraphs[0]
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run("KOA MARIE GERVAIS NELLY")
    run_name.font.name = 'Segoe UI Semibold'
    run_name.font.size = Pt(22)
    run_name.font.bold = True
    run_name.font.color.rgb = WHITE

    # Header Subtitle
    p_sub = cell.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(6)
    run_sub = p_sub.add_run("Lead AI Engineer & Consultant IA / Data | Fondateur @ Archi Cam AI")
    run_sub.font.size = Pt(11)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0x00, 0xF2, 0xFE)

    # Badge Lab
    p_badge = cell.add_paragraph()
    p_badge.paragraph_format.space_before = Pt(0)
    p_badge.paragraph_format.space_after = Pt(8)
    run_badge = p_badge.add_run("★ Candidature en cours - Google Africa Applied AI Lab (Accra, Ghana)")
    run_badge.font.size = Pt(9.5)
    run_badge.font.bold = True
    run_badge.font.color.rgb = WHITE

    # Contacts Line inside Header
    p_cnt = cell.add_paragraph()
    p_cnt.paragraph_format.space_before = Pt(4)
    p_cnt.paragraph_format.space_after = Pt(0)
    run_cnt = p_cnt.add_run("Email : koagervais85@gmail.com  |  Tél : +237 695 35 34 02  |  Douala / Ngaoundéré, Cameroun  |  GitHub : github.com/gervais-afk")
    run_cnt.font.size = Pt(8.5)
    run_cnt.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Helper function for Section Titles
    def add_section_header(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title_text.upper())
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = DARK_BLUE
        add_bottom_border(p, color_hex="00A8B5", size="12")
        return p

    # 2. PROFIL SYNTHÉTIQUE
    add_section_header("Profil Synthétique")
    p_prof = doc.add_paragraph()
    p_prof.paragraph_format.space_after = Pt(6)
    p_prof.paragraph_format.line_spacing = 1.15
    run_prof = p_prof.add_run(
        "Ingénieur & Spécialiste en IA Appliquée alliant la rigueur du Génie Civil et l'expertise en Sûreté Opérationnelle à la maîtrise des architectures émergentes d'IA Générative, GraphRAG (Neo4j) et MLOps. Fondateur d'Archi Cam AI (projet développé pour le Google Africa Applied AI Lab). Concepteur de solutions IA souveraines d'entreprise éliminant les hallucinations des LLM par l'intégration de moteurs déterministes en Sandbox."
    )
    run_prof.font.size = Pt(9.5)

    # 3. RÉALISATIONS & PROJETS IA MAJEURS
    add_section_header("Projets IA Majeurs & Réalisations")

    def add_project_item(title, category, bullets):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        
        r_t = p.add_run(title + " ")
        r_t.font.bold = True
        r_t.font.size = Pt(10)
        r_t.font.color.rgb = DARK_BLUE
        
        r_c = p.add_run("| " + category)
        r_c.font.bold = True
        r_c.font.size = Pt(9)
        r_c.font.color.rgb = CYAN

        for b in bullets:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(1.5)
            bp.paragraph_format.left_indent = Inches(0.25)
            brun = bp.add_run(b)
            brun.font.size = Pt(9)

    add_project_item(
        "1. Archi Cam AI 🏛️", 
        "SaaS IA Agentique & 5D BIM (Projet soumis au Google Africa Applied AI Lab)",
        [
            "Plateforme IA souveraine de modélisation BIM 5D et génération automatisée de devis BTP normés (BAEL 91 / Eurocodes).",
            "Combinaison de Google Gemma 4 12B local, Gemini 1.5 Pro Vision et d'un moteur Python Sandbox (IfcOpenShell, BAEL 91) pour des calculs sans hallucination.",
            "Génération de devis Excel (DQE) en moins de 2 minutes et rendus photoréalistes via Imagen 3 + ControlNet."
        ]
    )

    add_project_item(
        "2. Sovereign.BI Agentic 📊", 
        "Business Intelligence Agentique & Security",
        [
            "Moteur décisionnel autonome permettant d'interroger des bases de données SQL complexes en langage naturel (React, FastAPI, PostgreSQL).",
            "Architecture TypeScript Orchestrator et Neo4j N10S (GraphRAG).",
            "Intégration de guardrails dynamiques anti-injection et d'un auditeur d'explicabilité SHAP Sentinel."
        ]
    )

    add_project_item(
        "3. Dataset Automator ⚙️", 
        "MLOps & Traçabilité Séries Temporelles",
        [
            "Pipeline MLOps & RAG d'évaluation et de gouvernance de séries temporelles complexes.",
            "Indexation sémantique Neo4j, suivi des métriques en temps réel avec MLflow et synthèses explicatives générées par Firebase Genkit & Gemma-2."
        ]
    )

    add_project_item(
        "4. VigieSahel 🌾", 
        "IA Impact Climat & Santé Publique",
        [
            "Plateforme prédictive d'optimisation des dates de semis agricoles et d'anticipation des risques épidémiques/pollution PM2.5 (Streamlit, Supabase, ML, PWA Offline-First)."
        ]
    )

    # 4. COMPÉTENCES TECHNIQUES
    add_section_header("Compétences Techniques")
    
    def add_skill_row(category, skills):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run("• " + category + " : ")
        r1.font.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = DARK_BLUE
        
        r2 = p.add_run(skills)
        r2.font.size = Pt(9)

    add_skill_row("IA Générative & Agents", "Google Gemma 4, Gemini 1.5 Pro, Firebase Genkit, Agents Autonomes, RAG / GraphRAG, Prompt Engineering")
    add_skill_row("Data & Graphes", "Neo4j (Cypher, N10S), PostgreSQL, Supabase, MLflow (MLOps), Pandas, NumPy")
    add_skill_row("Dev & Génie Logiciel", "Next.js 14, React, FastAPI, Streamlit, Python (BAEL 91), IfcOpenShell (BIM 5D), Docker, Git & GitHub")

    # 5. PARCOURS PROFESSIONNEL
    add_section_header("Parcours Professionnel")

    def add_job_item(title, period, company, points):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        
        r_t = p.add_run(title)
        r_t.font.bold = True
        r_t.font.size = Pt(9.5)
        r_t.font.color.rgb = DARK_BLUE
        
        r_d = p.add_run(f"  ({period})")
        r_d.font.italic = True
        r_d.font.size = Pt(8.5)
        r_d.font.color.rgb = GRAY_TEXT
        
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_before = Pt(0)
        p_c.paragraph_format.space_after = Pt(2)
        r_comp = p_c.add_run(company)
        r_comp.font.bold = True
        r_comp.font.size = Pt(8.5)
        r_comp.font.color.rgb = CYAN

        for pt in points:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(1)
            bp.paragraph_format.left_indent = Inches(0.25)
            brun = bp.add_run(pt)
            brun.font.size = Pt(9)

    add_job_item(
        "Consultant IA & Data Science",
        "2025 - Présent",
        "Projets Indépendants & Entreprises | Douala",
        [
            "Analyse exploratoire et prétraitement de jeux de données massifs complexes.",
            "Modélisation de graphes de connaissances (Neo4j Cypher) et développement de pipelines RAG.",
            "Conception de bases de données SQL/PostgreSQL et reporting décisionnel interactif."
        ]
    )

    add_job_item(
        "Agent de Sûreté Aéroportuaire (AVSEC)",
        "2018 - Présent",
        "CCAA (Autorité Aéronautique du Cameroun)",
        [
            "Analyse des risques opérationnels critiques et contrôle strict des accès sécurisés.",
            "Rédaction de rapports d'audit de sûreté et coordination d'interventions d'urgence."
        ]
    )

    # 6. FORMATION ACADÉMIQUE
    add_section_header("Formation Académique")

    def add_edu_item(degree, period, school, desc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        r_d = p.add_run(degree)
        r_d.font.bold = True
        r_d.font.size = Pt(9.5)
        r_d.font.color.rgb = DARK_BLUE
        
        r_p = p.add_run(f"  ({period})")
        r_p.font.italic = True
        r_p.font.size = Pt(8.5)
        r_p.font.color.rgb = GRAY_TEXT

        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(0)
        p_s.paragraph_format.space_after = Pt(2)
        r_sc = p_s.add_run(school + " — " + desc)
        r_sc.font.size = Pt(8.5)

    add_edu_item("Master 2 Intelligence Artificielle & Data Science", "2025 - 2027", "Université de Ngaoundéré", "Spécialisation en Modélisation de Graphes (Neo4j), MLOps, Prompt Engineering & LLM.")
    add_edu_item("Licence & BTS Génie Civil (Option Bâtiment)", "2015 - 2016", "ISTDI / IUC Douala", "Dimensionnement de structures (BAEL 91), calculs de métrés et gestion de projets BTP.")

    # 7. LANGUES & ATOUTS
    add_section_header("Langues & Atouts Clés")
    p_end = doc.add_paragraph()
    p_end.paragraph_format.space_before = Pt(2)
    p_end.paragraph_format.space_after = Pt(0)
    r_l = p_end.add_run("• Langues : ")
    r_l.font.bold = True
    r_l.font.color.rgb = DARK_BLUE
    p_end.add_run("Français (Courant), Anglais (Technique & Professionnel)  |  ")
    r_a = p_end.add_run("• Atouts : ")
    r_a.font.bold = True
    r_a.font.color.rgb = DARK_BLUE
    p_end.add_run("Double compétence IA & Génie Civil, Rigueur de calcul & Sécurité des données.")

    output_path = r"c:\Users\HP\Desktop\portfolio-gervais\KOA_MARIE_GERVAIS_NELLY_CV.docx"
    doc.save(output_path)
    print(f"CV Word generated successfully at: {output_path}")

if __name__ == "__main__":
    generate_cv_docx()
